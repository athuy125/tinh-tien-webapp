import streamlit as st
import os
import json
import zipfile
from datetime import datetime
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2 import service_account
import pytz
from zoneinfo import ZoneInfo
import glob
import schedule
import time
from db import save_data, load_data, get_history, save_tinh_toan

# Tạo thư mục nếu chưa có
DATA_FOLDER = "data"
BACKUP_FOLDER = "backups"
os.makedirs(DATA_FOLDER, exist_ok=True)
os.makedirs(BACKUP_FOLDER, exist_ok=True)

# Hàm tạo file backup .zip
def backup_data_folder():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"backup_{timestamp}.zip"
    backup_path = os.path.join(BACKUP_FOLDER, backup_name)
    with zipfile.ZipFile(backup_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for foldername, _, filenames in os.walk(DATA_FOLDER):
            for filename in filenames:
                filepath = os.path.join(foldername, filename)
                arcname = os.path.relpath(filepath, DATA_FOLDER)
                zipf.write(filepath, arcname)
    return backup_path

# Hàm upload lên Google Drive
def upload_to_drive(local_file_path, drive_folder_id):
    SCOPES = ['https://www.googleapis.com/auth/drive']
    SERVICE_ACCOUNT_FILE = 'credentials.json'

    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=creds)

    file_metadata = {'name': os.path.basename(local_file_path), 'parents': [drive_folder_id]}
    media = MediaFileUpload(local_file_path, resumable=True)
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file.get("id")



# Hàm nạp và lưu dữ liệu
def load_data():
    if not os.path.exists(DATA_FILE):
        return {"debtors": [], "notes": [], "history": []}
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_data(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Hàm thêm dữ liệu
def add_note(content):
    data = load_data()
    data["notes"].append({"content": content, "time": datetime.now().isoformat()})
    save_data(data)



def add_history(action, detail):
    data = load_data()
    data["history"].append({"action": action, "detail": detail, "time": datetime.now().isoformat()})
    save_data(data)













# PWA header
st.markdown("""
<link rel="manifest" href="/manifest.json">
<script>
  if ('serviceWorker' in navigator) {
    navigator.serviceWorker.register('/service-worker.js')
      .then(function(reg) { console.log('✅ Service Worker đăng ký thành công:', reg); })
      .catch(function(err) { console.log('⚠️ Service Worker lỗi:', err); });
  }
</script>
""", unsafe_allow_html=True)

# Page config
st.set_page_config(page_title="📦 Công cụ Tính Tiền & Quản Lý Nợ by Huyhihihi", layout="centered")

# CSS nền tối
st.markdown("""
<style>
html, body, [class*="st-"], [class^="st-"] {
    background-color: #1e1e1e !important;
    color: #fdfdfd !important;
}
h1, h2, h3 {
    color: #fdfdfd !important;
}
label, .st-bb, .st-cx, .st-eb, .css-1cpxqw2 {
    color: #fdfdfd !important;
    font-size: 18px !important;
}
input, textarea {
    background-color: #333 !important;
    color: #fdfdfd !important;
}
.stButton>button {
    background-color: #444 !important;
    color: #fdfdfd !important;
    font-size: 18px !important;
}
</style>
""", unsafe_allow_html=True)

# Tiêu đề
st.title("📦 Công cụ Tính Tiền & Quản Lý Nợ by Huyhihihi")

username = st.text_input("👉 Nhập tên của bạn để bắt đầu:")


# ====== CẤU HÌNH ======



# Tạo thư mục nếu chưa có




def add_history(data, section, info):
    """
    Lưu lại lịch sử tính toán.
    section: "profit" hoặc "import"
    info: chuỗi mô tả
    """
    time = datetime.now(ZoneInfo("Asia/Ho_Chi_Minh")).strftime("%Y-%m-%d %H:%M:%S")
    history = data.get("history", {})
    if section not in history:
        history[section] = []
    history[section].append(f"{time}: {info}")
    data["history"] = history
    save_data(data)
def get_filename(username):
    return os.path.join(DATA_FOLDER, f"data_{username}.json")
# Hàm phục hồi



def load_data(username):
    filepath = get_filename(username)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_data(username, data):
    filepath = get_filename(username)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def log_action(data, action):
    logs = data.get("logs", [])
    time = datetime.now(ZoneInfo("Asia/Ho_Chi_Minh")).strftime("%Y-%m-%d %H:%M:%S")
    logs.append(f"{time}: {action}")
    data["logs"] = logs
    return data

# ====== SAO LƯU & PHỤC HỒI ======


def parse_sl(text):
    """
    Chấp nhận các biểu thức như: 11+5, 11*2, 10/2, 20-3
    """
    try:
        text = text.replace(',', '.')  # đổi , thành .
       
        allowed = "0123456789+-*/.() "
        if all(c in allowed for c in text):
            return eval(text)
        else:
            return 0
    except:
        return 0
if username:
    filename = f"data_{username}.json"

    def load_data():
        if os.path.exists(filename):
            with open(filename, "r", encoding="utf-8") as f:
                return json.load(f)
        return {}

    def save_data(data):
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)

    def log_action(action):
        logs = data.get("logs", [])
        time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        logs.append(f"{time}: {action}")
        data["logs"] = logs
        save_data(data)

    data = load_data()
    is_vip = data.get("is_vip", False)



    if is_vip:
        st.success(f"🌟 {username}, bạn đang là THÀNH VIÊN VIP! 🌟")
    

    # Menu
    menu = [
        "Tính lời + vốn (gộp)", 
        "Quản lý nợ", 
        "Tính thuế (VIP)", 
        "💼 Lợi nhuận chuyến xe đầu kéo",
        "🌟 Thông tin VIP & Thanh toán",
        "📊 Thống kê & Xuất dữ liệu (VIP)",
        "📝 Ghi chú cá nhân (VIP)",
        "📊 Máy tính phần trăm (VIP)",
        "📜 Nhật ký hoạt động (VIP)",
        "🛡 Sao lưu & Phục hồi dữ liệu",
        "📜 Lịch sử tính toán theo mặt hàng",

    ]

    choice = st.sidebar.selectbox("📌 Chọn chức năng", menu)
    st.markdown("<hr style='margin:20px 0'>", unsafe_allow_html=True)

    # Tính tiền lời
    if choice == "Tính lời + vốn (gộp)":
        st.subheader("💰 Tính lời + vốn (gộp)")
        ten_hang = st.text_input("Tên mặt hàng (ví dụ: sầu riêng, vải...)")
        sl_thu_text = st.text_input("Số lượng thu (nhập) VD: 11,5 hoặc 11+5")
        gia_thu = st.number_input("Giá thu / đơn vị (nghìn đồng)", 0, step=1)
        sl_ban_text = st.text_input("Số lượng bán VD: 10 hoặc 5+3")
        gia_ban = st.number_input("Giá bán / đơn vị (nghìn đồng)", 0, step=1)

        if st.button("✅ Tính"):
            sl_thu = parse_sl(sl_thu_text)
            sl_ban = parse_sl(sl_ban_text)

            tien_von = sl_thu * gia_thu
            tien_ban = sl_ban * gia_ban
            tien_loi = tien_ban - tien_von

            st.info(f"👉 Tổng số lượng thu (nhập): **{sl_thu}**")
            st.info(f"👉 Tiền vốn (nhập): **{tien_von} nghìn đồng**")
            st.info(f"👉 Tổng số lượng bán: **{sl_ban}**")
            st.info(f"👉 Tiền bán: **{tien_ban} nghìn đồng**")
            st.success(f"✅ Tiền lời: **{tien_loi} nghìn đồng**")
            


            # --- Lưu lịch sử theo mặt hàng ---
            if ten_hang.strip():
                noi_dung = f"{ten_hang}: thu {sl_thu}×{gia_thu}={tien_von}, bán {sl_ban}×{gia_ban}={tien_ban}, lời {tien_loi}"
                history = data.get("history", {})
                if ten_hang not in history:
                    history[ten_hang] = []
                history[ten_hang].append(noi_dung)
                data["history"] = history
                save_data(data)
                save_tinh_toan(username, ten_hang, noi_dung)
                st.success("✅ Đã lưu vào MongoDB!")
                
    # Lợi nhuận xe đầu kéo
    elif choice == "💼 Lợi nhuận chuyến xe đầu kéo":
        st.subheader("🚚 Tính lợi nhuận 1 chuyến xe đầu kéo")
        doanh_thu = st.number_input("💰 Doanh thu (triệu đồng)", 0.0, step=0.1)
        xang = st.number_input("⛽ Xăng dầu (triệu đồng)", 0.0, step=0.1)
        cau = st.number_input("🛣️ Phí cầu đường (triệu đồng)", 0.0, step=0.1)
        sua = st.number_input("🔧 Sửa chữa (triệu đồng)", 0.0, step=0.1)
        cuoc = st.number_input("🚚 Cước xe (triệu đồng)", 0.0, step=0.1)
        an = st.number_input("🍚 Ăn uống (triệu đồng)", 0.0, step=0.1)
        if st.button("✅ Tính lợi nhuận"):
            tong_cp = xang + cau + sua + cuoc + an
            loi_nhuan = doanh_thu - tong_cp
            st.info(f"👉 **Tổng chi phí:** {tong_cp:.2f} triệu")
            st.success(f"✅ **Lợi nhuận:** {loi_nhuan:.2f} triệu")

    # Quản lý nợ
    elif choice == "Quản lý nợ":
        st.subheader("📝 Quản lý danh sách nợ")
        list_no = {k:v for k,v in data.items() if k not in ["is_vip","vip_amount","logs","notes","history"]}

        if list_no:
            ten = st.selectbox("👉 Chọn người nợ:", list(list_no.keys()))
            if ten:
                st.write(f"**Số nợ hiện tại của {ten}:** {list_no[ten]}")
                action = st.radio("Chọn tác vụ", [
                    "➕ Cộng thêm nợ", "✅ Trả bớt nợ", "📦 Tính nợ theo số thùng",
                    "✏️ Đổi tên người nợ", "🗑️ Xóa người nợ"
                ])
                if action == "➕ Cộng thêm nợ":
                    them = st.number_input("Số tiền muốn cộng (nghìn đồng)", 0, step=1)
                    if st.button("Cộng thêm"):
                        try: cu = int(str(list_no[ten]).split()[0])
                        except: cu = 0
                        moi = cu + them
                        data[ten] = f"{moi} (Đã nợ {cu} + thêm {them})"
                        save_data(data)
                        st.success(f"✅ Tổng nợ mới: {moi} nghìn đồng")
                elif action == "✅ Trả bớt nợ":
                    tra = st.number_input("Số tiền muốn trả (nghìn đồng)", 0, step=1)
                    if st.button("Cập nhật sau khi trả"):
                        try: cu = int(str(list_no[ten]).split()[0])
                        except: cu = 0
                        moi = max(cu - tra, 0)
                        data[ten] = f"{moi} (Đã trả {tra} từ {cu})"
                        save_data(data)
                        st.success(f"✅ Nợ mới: {moi} nghìn đồng")
                elif action == "📦 Tính nợ theo số thùng":
                    sl = st.number_input("Số thùng nợ thêm", 0, step=1)
                    gia = st.number_input("Giá bán / thùng (nghìn đồng)", 0, step=1)
                    if st.button("Tính & Cập nhật"):
                        them = sl * gia
                        try: cu = int(str(list_no[ten]).split()[0])
                        except: cu = 0
                        moi = cu + them
                        data[ten] = f"{moi} (Đã nợ {cu} + thêm {them})"
                        save_data(data)
                        st.success(f"✅ Tổng nợ mới: {moi} nghìn đồng")
                elif action == "✏️ Đổi tên người nợ":
                    ten_moi = st.text_input("Nhập tên mới")
                    if st.button(f"Đổi tên {ten} → {ten_moi}"):
                        if ten_moi and ten_moi not in data:
                            data[ten_moi] = data.pop(ten)
                            save_data(data)
                            st.success(f"✅ Đã đổi tên {ten} thành {ten_moi}")
                        else:
                            st.error("⚠️ Tên mới bị trống hoặc đã tồn tại!")
                elif action == "🗑️ Xóa người nợ":
                    if st.button(f"Xóa {ten}"):
                        del data[ten]
                        save_data(data)
                        st.success("✅ Đã xóa người nợ")
        else:
            st.info("Danh sách nợ trống.")
        st.markdown("---")
        st.subheader("➕ Thêm người nợ mới")
        ten_moi = st.text_input("Tên người nợ mới")
        so_moi = st.text_input("Số tiền nợ")
        if st.button("Thêm"):
            data[ten_moi] = so_moi
            save_data(data)
            save_debt(username, ten_moi, so_moi)
            st.success(f"✅ Đã thêm: {ten_moi}")

    # VIP
    elif choice == "🌟 Thông tin VIP & Thanh toán":
        st.subheader("🌟 Đăng ký VIP (LƯU Ý: SAU KHI ĐĂNG KÝ GÓI VIP VUI LÒNG CHÚ Ý ĐIỆN THOẠI, SẼ CÓ NGƯỜI GỌI ĐẾN CUNG CẤP MÃ VIP CHO BẠN, GỌI SỐ 0937481127 nếu muốn liên hệ)")
        st.markdown("""
        **🏦 Ngân hàng:** Techcombank  
        **👤 Chủ tài khoản:** Đỗ Hoàng Gia Huy  
        **💳 Số tài khoản:** 7937481127  
        **💰 Nội dung:** VIP + [Your name] + [phone number]
        """)
        vip_amount = st.number_input("Số tiền đã chuyển (nghìn đồng)", 0, step=1)
        secret_code = st.text_input("Nhập mã bí mật bạn nhận được")
        if st.button("✅ Xác nhận VIP"):
            if secret_code == "521985":
                data["is_vip"] = True
                data["vip_amount"] = vip_amount
                save_data(data)
                st.success("🌟 Chúc mừng! Bạn đã trở thành VIP!")
            else:
                st.warning("⚠️ Mã không đúng, vui lòng kiểm tra.")

    # Tính thuế (VIP)
    if choice == "Tính thuế (VIP)":
        if is_vip:
            st.subheader("💵 Tính thuế")
            tab = st.radio("Chọn loại thuế", ["TNCN (tiền lương)", "Thuế kinh doanh", "Thuế bán hàng (GTGT)"])
            if tab == "TNCN (tiền lương)":
                luong = st.number_input("Tổng thu nhập (triệu đồng/tháng)", 0.0, step=0.1)
                phu_thuoc = st.number_input("Số người phụ thuộc", 0, step=1)
                hop_dong = st.checkbox("Hợp đồng ≥3 tháng?", value=True)
                if st.button("✅ Tính"):
                    giam_tru = 11 + phu_thuoc * 4.4
                    tntt = max(luong - giam_tru, 0)
                    if not hop_dong:
                        thue = luong * 0.10
                        pp = "Khấu trừ 10%"
                    else:
                        t = tntt
                        if t<=0: thue=0
                        elif t<=5: thue=0.05*t
                        elif t<=10: thue=0.10*t-0.25
                        elif t<=18: thue=0.15*t-0.75
                        elif t<=32: thue=0.20*t-1.65
                        elif t<=52: thue=0.25*t-3.25
                        elif t<=80: thue=0.30*t-5.85
                        else: thue=0.35*t-9.85
                        pp = "Biểu thuế lũy tiến"
                    con_lai= luong - thue
                    st.info(f"TNTT: {tntt:.2f} triệu ({pp})")
                    st.success(f"Thuế: {thue:.2f} triệu | Sau thuế: {con_lai:.2f} triệu")
            elif tab=="Thuế kinh doanh":
                dt = st.number_input("Doanh thu năm (triệu)",0.0,step=0.1)
                if st.button("Tính"):
                    thue_gtgt=dt*0.10; thue_tncn=dt*0.01
                    st.success(f"GTGT: {thue_gtgt:.2f} | TNCN: {thue_tncn:.2f} | Tổng: {thue_gtgt+thue_tncn:.2f}")
            else:
                dt = st.number_input("Doanh thu bán hàng (triệu)",0.0,step=0.1)
                if st.button("Tính"):
                    thue=dt*0.10
                    st.success(f"Thuế GTGT: {thue:.2f} triệu")
        else:
            st.warning("🌟 Vui lòng nâng cấp VIP để dùng tính năng này!")

    # Thống kê & Xuất dữ liệu (VIP)
    elif choice=="📊 Thống kê & Xuất dữ liệu (VIP)":
        if is_vip:
            list_no={k:v for k,v in data.items() if k not in ["is_vip","vip_amount","logs","notes"]}
            tong_no=sum(int(str(v).split()[0]) for v in list_no.values() if str(v).split()[0].isdigit())
            st.metric("Số người nợ",len(list_no))
            st.metric("Tổng nợ (nghìn)",tong_no)
            if st.button("📥 Xuất JSON"):
                st.download_button("Tải JSON", json.dumps(data,ensure_ascii=False,indent=4), file_name=f"{username}_data.json")
            if st.button("📥 Xuất Word"):
                doc=Document(); doc.add_heading(f"Dữ liệu của {username}",0)
                for k,v in data.items(): doc.add_paragraph(f"{k}: {v}")
                tmp="temp.docx"; doc.save(tmp)
                with open(tmp,"rb") as f:
                    st.download_button("Tải Word",f, file_name=f"{username}_data.docx")
        else:
            st.warning("🌟 Vui lòng nâng cấp VIP để dùng tính năng này!")

    # Ghi chú cá nhân (VIP)
    elif choice == "📝 Ghi chú cá nhân (VIP)":
        if is_vip:
            st.subheader("📝 Ghi chú cá nhân")
            # Lấy notes an toàn
            notes = data.get("notes", [])
            if not isinstance(notes, list):
                notes = []
                data["notes"] = notes
                save_data(data)

            new_note = st.text_area("Thêm ghi chú mới")

        if st.button("✅ Lưu ghi chú"):
            if new_note.strip():
                notes.append(new_note.strip())
                data["notes"] = notes
                save_data(data)
                st.success("Đã lưu ghi chú!")
                # Thêm log
                logs = data.get("logs", [])
                time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                logs.append(f"{time}: Thêm ghi chú: {new_note.strip()}")
                data["logs"] = logs
                save_data(data)

        st.markdown("---")

        if notes:
            st.subheader("📌 Danh sách ghi chú:")
            for i, note in enumerate(notes, 1):
                st.markdown(f"**{i}.** {note}")
                idx_xoa = st.number_input("Nhập số thứ tự ghi chú muốn xóa", min_value=1, max_value=len(notes), step=1)
                if st.button("🗑️ Xóa ghi chú"):
                    if 1 <= idx_xoa <= len(notes):
                        removed = notes.pop(idx_xoa-1)
                        data["notes"] = notes
                        save_data(data)
                        # Thêm log
                        logs = data.get("logs", [])
                        time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        logs.append(f"{time}: Xóa ghi chú: {removed}")
                        data["logs"] = logs
                        save_data(data)
                        st.success(f"Đã xóa: {removed}")
            else:
                st.info("Chưa có ghi chú nào.")
        else:
            st.warning("🌟 Vui lòng nâng cấp VIP để dùng tính năng này!")
    # Máy tính phần trăm (VIP)
    elif choice == "📊 Máy tính phần trăm (VIP)":
        if is_vip:
            st.subheader("📊 Máy tính phần trăm")
            so_goc = st.number_input("Giá trị gốc", 0.0, step=0.1)
            phan_tram = st.number_input("Tỷ lệ phần trăm (%)", 0.0, step=0.1)
            phep = st.radio("Chọn phép tính", ["Tăng thêm", "Giảm bớt"])
            if st.button("✅ Tính"):
                ket_qua = so_goc * (1 + phan_tram/100) if phep=="Tăng thêm" else so_goc * (1 - phan_tram/100)
                st.success(f"Kết quả: {ket_qua:.2f}")
                log_action(f"Tính phần trăm: {phep} {phan_tram}% của {so_goc} = {ket_qua}")
        else:
            st.warning("🌟 Vui lòng nâng cấp VIP để dùng tính năng này!")

    # Nhật ký hoạt động (VIP)
    elif choice == "📜 Nhật ký hoạt động (VIP)":
        if is_vip:
            st.subheader("📜 Nhật ký hoạt động")
            logs = data.get("logs", [])
            if logs:
                for log in reversed(logs[-50:]):
                    st.markdown(f"- {log}")
            else:
                st.info("Chưa có hoạt động nào.")
        else:
            st.warning("🌟 Vui lòng nâng cấp VIP để dùng tính năng này!")
    elif choice == "🛡 Sao lưu & Phục hồi dữ liệu":
        st.title("🛡 Tự động backup & upload Google Drive")

        if st.button("📦 Sao lưu và Upload"):
            backup_file = backup_data_folder()
            st.success(f"✅ Đã tạo file backup: {os.path.basename(backup_file)}")

            try:
                drive_folder_id = "1TLcveIa9xgbgOLXfCnR48_fLAh1uVhPj"  # <-- Thay bằng folder ID ở Bước 4
                file_id = upload_to_drive(backup_file, drive_folder_id)
                st.success(f"📤 Đã upload lên Google Drive! File ID: {file_id}")
            except Exception as e:
                st.error(f"❌ Lỗi upload: {e}")

    elif choice == "📜 Lịch sử tính toán theo mặt hàng":
        st.subheader("📜 Lịch sử tính toán theo mặt hàng")
        

        history = data.get("history", {})
        list_mat_hang = list(history.keys())
        st.markdown("---")
        st.subheader("🧮 Tính toán từ dữ liệu lịch sử")

        cong_thuc = st.text_input("✏️ Nhập công thức (ví dụ: 893432514 + 10000 * 2):")
        profit_history = []
        if st.button("✅ Tính"):
            try:
                if cong_thuc.strip():  # kiểm tra không rỗng
                    ket_qua = eval(cong_thuc, {"__builtins__": {}})
                    st.success(f"📌 Kết quả: **{ket_qua}**")
                    # Lưu
                    new_line = f"Tổng tiền của {cong_thuc} = {ket_qua}"
                    profit_history.append(new_line)
                    history["profit"] = profit_history
                    data["history"] = history
                    save_data(data)
                    save_tinh_toan(username, hang, content)
                    st.info("✅ Đã lưu vào lịch sử tính toán!")
                else:
                    st.warning("⚠️ Vui lòng nhập công thức trước khi tính.")
            except Exception as e:
                    st.error(f"❌ Lỗi: {e}")


        if list_mat_hang:
            
            
            selected_hang = st.selectbox("📌 Chọn mặt hàng để xem lịch sử", list(history.keys()))
            if selected_hang:
                # Lấy từ MongoDB
                online_history = get_history(username, selected_hang)
                st.markdown(f"### 🧾 Lịch sử online của **{selected_hang}**:")
                for item in online_history:
                    st.markdown(f"- {item['content']}")
                st.markdown(f"### 🧾 Lịch sử của **{selected_hang}**:")
                items = history.get(selected_hang, [])
                for i, item in enumerate(reversed(history[selected_hang]), 1):
                    st.markdown(f"**{i}.** {item}")
                    

                # Thêm nút xoá lịch sử từng dòng
                if len(items) > 0:
                    idx_xoa = st.number_input(
                        "Nhập số thứ tự dòng muốn xoá",
                        min_value=1,
                        max_value=len(items),
                        step=1,
                        key=f"xoa_{selected_hang}_{username}"
                     )

                    if st.button("🗑️ Xoá dòng này"):
                        real_idx = len(items) - idx_xoa
                        removed = items.pop(real_idx)
                        history[selected_hang] = items
                        data["history"] = history
                        save_data(data)
                        st.success(f"✅ Đã xoá: {removed}")
                else:
                     st.info("⚠️ Chưa có lịch sử nào để xoá.")

        

          

else:
    st.info("👉 Nhập tên để bắt đầu.")




